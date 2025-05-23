def extract_text_from_docx(docx_file):
    """Ekstrak seluruh teks dari file docx sebagai satu string."""
    from docx import Document
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_paragraphs_from_docx(docx_file):
    """Ekstrak list paragraf (string) dari file docx."""
    from docx import Document
    doc = Document(docx_file)
    return [para.text for para in doc.paragraphs if para.text.strip()]

def extract_text_by_section(docx_file):
    """
    Ekstrak teks per section (jika ingin dikembangkan, misal per halaman/section).
    Saat ini, kembalikan dict {1: all_text} agar kompatibel dengan pages_text PDF.
    """
    text = extract_text_from_docx(docx_file)
    return {1: text}
